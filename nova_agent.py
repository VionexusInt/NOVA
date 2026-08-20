import os
import sys
import json
import subprocess
import webbrowser
import pyautogui
import pyperclip
import time
import threading
import hashlib
import platform
import shutil
import glob
import signal
import urllib.request
import xml.etree.ElementTree as ET
import difflib
import uuid
import warnings
import logging
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify
from flask_cors import CORS

warnings.filterwarnings('ignore')
logging.disable(logging.WARNING)

try:
    from winotify import Notification
    NOTIFICACIONES_WINDOWS_OK = True
    _notif_engine = 'winotify'
except ImportError:
    NOTIFICACIONES_WINDOWS_OK = False
    _notif_engine = None

if not NOTIFICACIONES_WINDOWS_OK:
    try:
        from win10toast_click import ToastNotifier
        _toaster = ToastNotifier()
        NOTIFICACIONES_WINDOWS_OK = True
        _notif_engine = 'win10toast'
    except ImportError:
        _toaster = None
        NOTIFICACIONES_WINDOWS_OK = False

app = Flask(__name__)
CORS(app)
pyautogui.FAILSAFE = True
pyautogui.PAUSE = 0.2

HOME = Path.home()
DESKTOP = HOME / "Desktop"
DOCUMENTS = HOME / "Documents"
DOWNLOADS = HOME / "Downloads"

NOVA_ROOT = Path(__file__).parent.resolve()

ARCHIVOS_AUTOMEJORA = [
    'js/agent.js', 'js/api.js', 'js/audio.js', 'js/briefing.js', 'js/calendar.js',
    'js/chat.js', 'js/email.js', 'js/helpers.js', 'js/init.js',
    'js/marketing.js', 'js/mejora.js', 'js/memoria.js', 'js/mic.js', 'js/orb.js',
    'js/paneles.js', 'js/programacion.js', 'js/state.js', 'js/tareas.js',
    'js/vision.js', 'js/wake.js', 'css/estilos.css', 'index.html',
    'nova_agent.py', 'piper_server.py',
]

propuestas_pendientes = {}

BACKUPS_DIR = NOVA_ROOT / '_backups_automejora'
BACKUPS_DIR.mkdir(exist_ok=True)

def _resolver_ruta_archivo(archivo):
    if not archivo or not isinstance(archivo, str):
        return None
    archivo = archivo.strip().strip('"\'`').replace('\\', '/')
    if archivo.startswith('./'):
        archivo = archivo[2:]
    if archivo not in ARCHIVOS_AUTOMEJORA:
        return None
    ruta = (NOVA_ROOT / archivo).resolve()
    try:
        ruta.relative_to(NOVA_ROOT)
    except ValueError:
        return None
    return ruta

def _generar_diff(original, nuevo, archivo):
    try:
        orig_lines = original.splitlines(keepends=True)
        new_lines = nuevo.splitlines(keepends=True)
        diff = difflib.unified_diff(
            orig_lines, new_lines,
            fromfile=f'a/{archivo}', tofile=f'b/{archivo}',
            lineterm='', n=3
        )
        return ''.join(diff)
    except Exception:
        return ''

def _crear_backup(ruta_archivo, archivo_nombre):
    try:
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_nombre = f"{ts}__{archivo_nombre.replace('/', '__')}"
        backup_path = BACKUPS_DIR / backup_nombre
        shutil.copy2(ruta_archivo, backup_path)
        archivos_backup = sorted(BACKUPS_DIR.glob(f"*__{archivo_nombre.replace('/', '__')}"))
        for viejo in archivos_backup[:-10]:
            try:
                viejo.unlink()
            except Exception:
                pass
        return str(backup_path)
    except Exception as e:
        return None

@app.route('/api/mejora/leer', methods=['POST'])
def mejora_leer():
    try:
        data = request.get_json() or {}
        archivo = data.get('archivo', '')
        ruta = _resolver_ruta_archivo(archivo)
        if not ruta:
            return jsonify({'ok': False, 'error': f'Archivo no permitido o inexistente: {archivo}'}), 400
        if not ruta.exists():
            return jsonify({'ok': False, 'error': f'El archivo no existe: {ruta}'}), 404
        contenido = ruta.read_text(encoding='utf-8')
        return jsonify({'ok': True, 'archivo': archivo, 'contenido': contenido})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/mejora/proponer', methods=['POST'])
def mejora_proponer():
    try:
        data = request.get_json() or {}
        archivo = data.get('archivo', '')
        contenido_nuevo = data.get('contenido_nuevo', '')
        descripcion = data.get('descripcion', 'Mejora de código')
        ruta = _resolver_ruta_archivo(archivo)
        if not ruta:
            return jsonify({'ok': False, 'error': f'Archivo no permitido: {archivo}'}), 400
        if not ruta.exists():
            return jsonify({'ok': False, 'error': f'El archivo no existe: {ruta}'}), 404
        if not contenido_nuevo or not contenido_nuevo.strip():
            return jsonify({'ok': False, 'error': 'Contenido nuevo vacío'}), 400
        contenido_actual = ruta.read_text(encoding='utf-8')
        if contenido_actual.strip() == contenido_nuevo.strip():
            return jsonify({'ok': True, 'lineas_cambiadas': 0, 'mensaje': 'Sin cambios reales'})
        diff = _generar_diff(contenido_actual, contenido_nuevo, archivo)
        lineas_cambiadas = diff.count('\n+') + diff.count('\n-')
        lineas_cambiadas = max(lineas_cambiadas, 1) if diff else 1
        propuesta_id = str(uuid.uuid4())[:8]
        propuestas_pendientes[propuesta_id] = {
            'archivo': archivo,
            'ruta': str(ruta),
            'contenido_nuevo': contenido_nuevo,
            'descripcion': descripcion,
            'timestamp': datetime.now().isoformat(),
            'diff': diff,
        }
        return jsonify({
            'ok': True,
            'propuesta_id': propuesta_id,
            'archivo': archivo,
            'descripcion': descripcion,
            'lineas_cambiadas': lineas_cambiadas,
            'diff': diff[:5000],
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/mejora/aprobar', methods=['POST'])
def mejora_aprobar():
    try:
        data = request.get_json() or {}
        propuesta_id = data.get('propuesta_id', '')
        if propuesta_id not in propuestas_pendientes:
            return jsonify({'ok': False, 'error': f'Propuesta {propuesta_id} no encontrada o expirada'}), 404
        propuesta = propuestas_pendientes[propuesta_id]
        archivo = propuesta['archivo']
        ruta = Path(propuesta['ruta'])
        if not ruta.exists():
            return jsonify({'ok': False, 'error': f'El archivo ya no existe: {ruta}'}), 404
        backup_path = _crear_backup(ruta, archivo)
        try:
            ruta.write_text(propuesta['contenido_nuevo'], encoding='utf-8')
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Error escribiendo archivo: {e}'}), 500
        del propuestas_pendientes[propuesta_id]
        notificar_windows('NOVA — Mejora aplicada', f'Archivo actualizado: {archivo}. Reinicia NOVA para ver los cambios.')
        return jsonify({
            'ok': True,
            'archivo': archivo,
            'backup': backup_path,
            'mensaje': f'Mejora aplicada en {archivo}. Reinicia NOVA para que surta efecto.',
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/mejora/rechazar', methods=['POST'])
def mejora_rechazar():
    try:
        data = request.get_json() or {}
        propuesta_id = data.get('propuesta_id', '')
        if propuesta_id in propuestas_pendientes:
            archivo = propuestas_pendientes[propuesta_id]['archivo']
            del propuestas_pendientes[propuesta_id]
            return jsonify({'ok': True, 'archivo': archivo, 'mensaje': 'Propuesta descartada'})
        return jsonify({'ok': True, 'mensaje': 'Propuesta no encontrada'})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/mejora/revertir', methods=['POST'])
def mejora_revertir():
    try:
        data = request.get_json() or {}
        archivo = data.get('archivo', '')
        ruta = _resolver_ruta_archivo(archivo)
        if not ruta:
            return jsonify({'ok': False, 'error': f'Archivo no permitido: {archivo}'}), 400
        sufijo = archivo.replace('/', '__')
        backups = sorted(BACKUPS_DIR.glob(f"*__{sufijo}"), reverse=True)
        if not backups:
            return jsonify({'ok': False, 'error': f'No hay backups disponibles de {archivo}'}), 404
        ultimo_backup = backups[0]
        try:
            shutil.copy2(ultimo_backup, ruta)
            return jsonify({
                'ok': True,
                'archivo': archivo,
                'backup_usado': ultimo_backup.name,
                'mensaje': f'{archivo} restaurado desde backup. Reinicia NOVA.',
            })
        except Exception as e:
            return jsonify({'ok': False, 'error': f'Error restaurando: {e}'}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/mejora/lista_archivos', methods=['GET'])
def mejora_lista_archivos():
    existentes = []
    for a in ARCHIVOS_AUTOMEJORA:
        ruta = NOVA_ROOT / a
        if ruta.exists():
            existentes.append({
                'archivo': a,
                'tamano': ruta.stat().st_size,
                'modificado': datetime.fromtimestamp(ruta.stat().st_mtime).isoformat()
            })
    backups = list(BACKUPS_DIR.iterdir()) if BACKUPS_DIR.exists() else []
    return jsonify({
        'ok': True,
        'archivos': existentes,
        'total_backups': len(backups),
        'propuestas_pendientes': len(propuestas_pendientes),
    })

def get_system_info():
    try:
        import psutil
        cpu = psutil.cpu_percent(interval=0.5)
        ram = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        battery = psutil.sensors_battery() if hasattr(psutil, 'sensors_battery') else None
        procs = sorted(
            [p.info for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent'])
            if p.info['cpu_percent'] > 0],
            key=lambda x: x['cpu_percent'], reverse=True
        )[:5]
        return {
            'cpu': round(cpu, 1),
            'ram_usada': round(ram.used/1e9, 1),
            'ram_total': round(ram.total/1e9, 1),
            'ram_pct': ram.percent,
            'disco_libre': round(disk.free/1e9, 1),
            'disco_total': round(disk.total/1e9, 1),
            'bateria': round(battery.percent) if battery else None,
            'cargando': battery.power_plugged if battery else None,
            'procesos_top': procs,
        }
    except ImportError:
        return {'error': 'psutil no instalado. Ejecuta: pip install psutil'}
    except Exception as e:
        return {'error': str(e)}

def get_procesos():
    try:
        import psutil
        procs = []
        for p in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'status']):
            try:
                procs.append(p.info)
            except:
                pass
        return sorted(procs, key=lambda x: x.get('cpu_percent', 0), reverse=True)[:20]
    except ImportError:
        return []

def matar_proceso(nombre_o_pid):
    try:
        import psutil
        killed = []
        for p in psutil.process_iter(['pid', 'name']):
            try:
                if str(p.info['pid']) == str(nombre_o_pid) or \
                    nombre_o_pid.lower() in p.info['name'].lower():
                    p.kill()
                    killed.append(p.info['name'])
            except:
                pass
        return f"Procesos terminados: {', '.join(killed)}" if killed else "No se encontró el proceso"
    except ImportError:
        return "psutil no instalado"

def listar_carpeta(ruta=None):
    if not ruta:
        ruta = DESKTOP
    ruta = Path(ruta)
    if not ruta.exists():
        return f"No existe: {ruta}"
    items = list(ruta.iterdir())
    return {
        'carpeta': str(ruta),
        'archivos': [{'nombre': i.name, 'tipo': 'carpeta' if i.is_dir() else 'archivo', 'tamaño': i.stat().st_size if i.is_file() else 0} for i in items[:30]]
    }

def buscar_archivo(nombre, donde=None):
    donde = donde or str(HOME)
    resultados = []
    try:
        for r in Path(donde).rglob(f"*{nombre}*"):
            resultados.append(str(r))
            if len(resultados) >= 20:
                break
    except:
        pass
    return resultados

def crear_carpeta(ruta):
    Path(ruta).mkdir(parents=True, exist_ok=True)
    return f"Carpeta creada: {ruta}"

def mover_archivo(origen, destino):
    shutil.move(origen, destino)
    return f"Movido: {origen} → {destino}"

def copiar_archivo(origen, destino):
    shutil.copy2(origen, destino)
    return f"Copiado: {origen} → {destino}"

def renombrar_archivo(origen, nuevo_nombre):
    p = Path(origen)
    nuevo = p.parent / nuevo_nombre
    p.rename(nuevo)
    return f"Renombrado: {p.name} → {nuevo_nombre}"

def abrir_archivo(ruta):
    subprocess.Popen(f'start "" "{ruta}"', shell=True)
    return f"Abriendo: {ruta}"

def abrir_carpeta(ruta=None):
    if not ruta:
        ruta = DESKTOP
    subprocess.Popen(f'explorer "{ruta}"', shell=True)
    return f"Abriendo carpeta: {ruta}"

def listar_ventanas():
    try:
        import pygetwindow as gw
        ventanas = gw.getAllTitles()
        return [v for v in ventanas if v.strip()]
    except:
        return []

def enfocar_ventana(titulo):
    try:
        import pygetwindow as gw
        wins = gw.getWindowsWithTitle(titulo)
        if wins:
            wins[0].activate()
            return f"Ventana enfocada: {titulo}"
        return f"No encontré ventana: {titulo}"
    except Exception as e:
        return f"Error: {e}"

def cerrar_ventana_titulo(titulo):
    try:
        import pygetwindow as gw
        wins = gw.getWindowsWithTitle(titulo)
        if wins:
            wins[0].close()
            return f"Ventana cerrada: {titulo}"
        return f"No encontré: {titulo}"
    except Exception as e:
        return f"Error: {e}"

def minimizar_ventana(titulo=None):
    if titulo:
        try:
            import pygetwindow as gw
            wins = gw.getWindowsWithTitle(titulo)
            if wins:
                wins[0].minimize()
                return f"Minimizada: {titulo}"
        except:
            pass
    pyautogui.hotkey('win', 'd')
    return "Escritorio mostrado"

def maximizar_ventana(titulo=None):
    if titulo:
        try:
            import pygetwindow as gw
            wins = gw.getWindowsWithTitle(titulo)
            if wins:
                wins[0].maximize()
                return f"Maximizada: {titulo}"
        except:
            pass
    pyautogui.hotkey('win', 'up')
    return "Ventana maximizada"

def abrir_programa(nombre):
    programas = {
        'chrome': 'chrome', 'google chrome': 'chrome',
        'firefox': 'firefox',
        'word': 'winword', 'microsoft word': 'winword',
        'excel': 'excel', 'microsoft excel': 'excel',
        'powerpoint': 'powerpnt',
        'notepad': 'notepad', 'bloc de notas': 'notepad',
        'calculadora': 'calc', 'calculator': 'calc',
        'explorador': 'explorer', 'file explorer': 'explorer',
        'vscode': 'code', 'visual studio code': 'code',
        'spotify': 'spotify', 'whatsapp': 'whatsapp',
        'teams': 'teams', 'zoom': 'zoom', 'discord': 'discord',
        'paint': 'mspaint', 'cmd': 'cmd', 'terminal': 'cmd',
        'powershell': 'powershell',
        'task manager': 'taskmgr', 'administrador de tareas': 'taskmgr',
        'configuracion': 'ms-settings:', 'settings': 'ms-settings:',
        'panel de control': 'control',
    }
    cmd = programas.get(nombre.lower(), nombre)
    subprocess.Popen(cmd, shell=True)
    return f"Abriendo {nombre}"

def abrir_web(url):
    if not url.startswith('http'):
        url = 'https://' + url
    webbrowser.open(url)
    return f"Abriendo {url}"

def buscar_google(query):
    url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
    webbrowser.open(url)
    return f"Buscando '{query}' en Google"

def escribir_texto(texto, delay=0.03):
    time.sleep(0.5)
    pyautogui.write(texto, interval=delay)
    return f"Texto escrito"

def tecla(key):
    pyautogui.press(key)
    return f"Tecla: {key}"

def hotkey(*keys):
    pyautogui.hotkey(*keys)
    return f"Hotkey: {'+'.join(keys)}"

def screenshot(nombre="captura"):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    ruta = DESKTOP / f"{nombre}_{ts}.png"
    pyautogui.screenshot(str(ruta))
    subprocess.Popen(f'start "" "{ruta}"', shell=True)
    return f"Captura guardada: {ruta.name}"

def portapapeles(texto):
    pyperclip.copy(texto)
    return "Copiado al portapapeles"

def volumen(accion):
    mapa = {'subir': ('volumeup', 5), 'bajar': ('volumedown', 5), 'silencio': ('volumemute', 1)}
    if accion in mapa:
        k, n = mapa[accion]
        for _ in range(n): pyautogui.press(k)
        return f"Volumen: {accion}"
    return "Acción desconocida"

def crear_word(nombre, contenido):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        t = doc.add_heading(nombre, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if isinstance(contenido, list):
            for item in contenido:
                tipo = item.get('tipo', 'parrafo')
                if tipo == 'titulo':
                    doc.add_heading(item['texto'], level=item.get('nivel', 1))
                elif tipo == 'parrafo':
                    doc.add_paragraph(item['texto'])
                elif tipo == 'lista':
                    for p in item.get('items', []):
                        doc.add_paragraph(p, style='List Bullet')
                elif tipo == 'tabla':
                    filas = item.get('filas', [])
                    if filas:
                        tbl = doc.add_table(rows=len(filas), cols=len(filas[0]))
                        tbl.style = 'Table Grid'
                        for i, fila in enumerate(filas):
                            for j, celda in enumerate(fila):
                                tbl.rows[i].cells[j].text = str(celda)
        else:
            for linea in str(contenido).split('\n'):
                doc.add_paragraph(linea)
        ruta = DESKTOP / f"{nombre}.docx"
        doc.save(str(ruta))
        subprocess.Popen(f'start "" "{ruta}"', shell=True)
        return f"Word creado: {nombre}.docx"
    except Exception as e:
        return f"Error Word: {e}"

def crear_pdf(nombre, contenido):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        ruta = DESKTOP / f"{nombre}.pdf"
        doc = SimpleDocTemplate(str(ruta), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = [Paragraph(nombre, styles['Title']), Spacer(1, 12)]
        if isinstance(contenido, list):
            for item in contenido:
                tipo = item.get('tipo', 'parrafo')
                if tipo == 'titulo':
                    story.append(Paragraph(item['texto'], styles['Heading1']))
                elif tipo == 'parrafo':
                    story.append(Paragraph(item['texto'], styles['Normal']))
                elif tipo == 'lista':
                    for p in item.get('items', []):
                        story.append(Paragraph(f"• {p}", styles['Normal']))
                    story.append(Spacer(1, 8))
        else:
            for linea in str(contenido).split('\n'):
                if linea.strip():
                    story.append(Paragraph(linea, styles['Normal']))
                story.append(Spacer(1, 4))
        doc.build(story)
        subprocess.Popen(f'start "" "{ruta}"', shell=True)
        return f"PDF creado: {nombre}.pdf"
    except Exception as e:
        return f"Error PDF: {e}"

def crear_imagen(nombre, texto, ancho=1920, alto=1080, bg='#000810', color='#00d4ff'):
    try:
        from PIL import Image, ImageDraw, ImageFont
        import textwrap
        img = Image.new('RGB', (int(ancho), int(alto)), color=bg)
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", int(alto*0.04))
            font_big = ImageFont.truetype("arialbd.ttf", int(alto*0.07))
        except:
            font = ImageFont.load_default()
            font_big = font
        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        lines = textwrap.wrap(texto, width=35)
        y = alto//2 - len(lines) * int(alto*0.05)
        for i, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font_big if i==0 else font)
            w = bbox[2] - bbox[0]
            draw.text(((ancho-w) //2, y), line, fill=(r, g, b), font=font_big if i==0 else font)
            y += int(alto*0.08) if i==0 else int(alto*0.055)
        ruta = DESKTOP / f"{nombre}.png"
        img.save(str(ruta))
        subprocess.Popen(f'start "" "{ruta}"', shell=True)
        return f"Imagen creada: {nombre}.png"
    except Exception as e:
        return f"Error imagen: {e}"

def ejecutar_cmd(cmd):
    try:
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=30)
        return result.stdout or result.stderr
    except Exception as e:
        return str(e)

automatizaciones = {}

def crear_automatizacion(nombre, condicion_tipo, condicion_valor, accion_tipo, accion_valor):
    automatizaciones[nombre] = {
        'condicion_tipo': condicion_tipo,
        'condicion_valor': condicion_valor,
        'accion_tipo': accion_tipo,
        'accion_valor': accion_valor,
        'activa': True,
        'creada': datetime.now().isoformat()
    }
    return f"Automatización creada: '{nombre}'"

def listar_automatizaciones():
    if not automatizaciones:
        return "Sin automatizaciones configuradas"
    return list(automatizaciones.keys())

def eliminar_automatizacion(nombre):
    if nombre in automatizaciones:
        del automatizaciones[nombre]
        return f"Eliminada: {nombre}"
    return f"No encontrada: {nombre}"

def monitor_automatizaciones():
    while True:
        try:
            import psutil
            for nombre, auto in list(automatizaciones.items()):
                if not auto['activa']:
                    continue
                tipo = auto['condicion_tipo']
                valor = auto['condicion_valor']
                disparar = False
                if tipo == 'cpu_mayor':
                    if psutil.cpu_percent() > float(valor):
                        disparar = True
                elif tipo == 'ram_mayor':
                    if psutil.virtual_memory().percent > float(valor):
                        disparar = True
                elif tipo == 'hora':
                    if datetime.now().strftime('%H:%M') == valor:
                        disparar = True
                if disparar:
                    accion = auto['accion_tipo']
                    aval = auto['accion_valor']
                    if accion == 'abrir_web':
                        abrir_web(aval)
                    elif accion == 'abrir_programa':
                        abrir_programa(aval)
                    elif accion == 'notificar':
                        notificar_frontend(nombre, aval)
                    auto['activa'] = False
        except:
            pass
        time.sleep(10)

notificaciones_pendientes = []

def notificar_frontend(titulo, mensaje, tambien_windows=True):
    notificaciones_pendientes.append({
        'titulo': titulo,
        'mensaje': mensaje,
        'ts': datetime.now().isoformat()
    })
    if tambien_windows:
        notificar_windows(titulo, mensaje)

def notificar_windows(titulo, mensaje, duracion=8):
    if not NOTIFICACIONES_WINDOWS_OK:
        return False
    def _mostrar():
        try:
            if _notif_engine == 'winotify':
                toast = Notification(
                    app_id="NOVA",
                    title=titulo[:64],
                    msg=mensaje[:200],
                    duration="short" if duracion <= 5 else "long",
                    icon=None
                )
                toast.show()
            elif _notif_engine == 'win10toast':
                try:
                    _toaster.show_toast(
                        titulo[:64],
                        mensaje[:200],
                        duration=duracion,
                        threaded=True,
                        icon_path=None
                    )
                except Exception:
                    pass
        except Exception:
            pass
    threading.Thread(target=_mostrar, daemon=True).start()
    return True

def monitor_proactivo():
    UMBRAL_CPU = 85
    UMBRAL_RAM = 90
    ultimo_aviso_cpu = 0
    ultimo_aviso_ram = 0
    while True:
        try:
            import psutil
            cpu = psutil.cpu_percent(interval=2)
            ram = psutil.virtual_memory().percent
            now = time.time()
            if cpu > UMBRAL_CPU and now - ultimo_aviso_cpu > 120:
                notificar_frontend('⚠ CPU Alta', f'CPU al {cpu}%. Revisar procesos.')
                ultimo_aviso_cpu = now
            if ram > UMBRAL_RAM and now - ultimo_aviso_ram > 120:
                notificar_frontend('⚠ RAM Alta', f'RAM al {ram}%. Sistema bajo presión.')
                ultimo_aviso_ram = now
            battery = psutil.sensors_battery()
            if battery and not battery.power_plugged and battery.percent < 15:
                notificar_frontend('🔋 Batería baja', f'Batería al {battery.percent}%. Conecta el cargador.')
        except:
            pass
        time.sleep(15)

@app.route('/api/ping', methods=['GET'])
def ping():
    return jsonify({'ok': True, 'mensaje': 'Agente NOVA v2 activo'})

@app.route('/api/sistema', methods=['GET'])
def sistema():
    return jsonify(get_system_info())

@app.route('/api/noticias', methods=['GET'])
def noticias():
    try:
        feeds = [
            'https://news.google.com/rss/search?q=Elche+Alicante&hl=es&gl=ES&ceid=ES:es',
            'https://news.google.com/rss/search?q=Elche&hl=es&gl=ES&ceid=ES:es',
            'https://news.google.com/rss/search?q=Alicante&hl=es&gl=ES&ceid=ES:es',
            'https://news.google.com/rss?hl=es&gl=ES&ceid=ES:es',
        ]
        titulos = []
        for feed_url in feeds:
            try:
                req = urllib.request.Request(feed_url, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                })
                with urllib.request.urlopen(req, timeout=8) as response:
                    xml_content = response.read().decode('utf-8')
                root = ET.fromstring(xml_content)
                for item in root.findall('.//item'):
                    title = item.find('title')
                    if title is not None and title.text:
                        titulos.append(title.text.strip())
                        if len(titulos) >= 8:
                            break
                if len(titulos) == 0:
                    for entry in root.findall('.//{http://www.w3.org/2005/Atom}entry'):
                        title = entry.find('{http://www.w3.org/2005/Atom}title')
                        if title is not None and title.text:
                            titulos.append(title.text.strip())
                            if len(titulos) >= 8:
                                break
                if len(titulos) > 0:
                    break
            except Exception:
                continue
        if len(titulos) == 0:
            titulos = ['Sistema de noticias temporalmente no disponible']
        return jsonify({'ok': True, 'titulos': titulos[:8]})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e), 'titulos': []}), 200

@app.route('/api/notificaciones', methods=['GET'])
def notificaciones():
    global notificaciones_pendientes
    ns = notificaciones_pendientes.copy()
    notificaciones_pendientes = []
    return jsonify(ns)

@app.route('/api/notificar_windows', methods=['POST'])
def notificar_windows_endpoint():
    try:
        data = request.get_json() or {}
        titulo = data.get('titulo', 'NOVA')
        mensaje = data.get('mensaje', '')
        if not mensaje:
            return jsonify({'ok': False, 'error': 'Falta el mensaje'}), 400
        if not NOTIFICACIONES_WINDOWS_OK:
            return jsonify({'ok': False, 'error': 'Sistema de notificaciones no disponible'}), 501
        ok = notificar_windows(titulo, mensaje)
        return jsonify({'ok': ok})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/notificaciones_estado', methods=['GET'])
def notificaciones_estado():
    return jsonify({'windows_disponible': NOTIFICACIONES_WINDOWS_OK, 'motor': _notif_engine})

@app.route('/api/accion', methods=['POST'])
def accion():
    data = request.get_json() or {}
    acc = data.get('accion', '')
    p = data.get('params', {})
    mapa = {
        'sistema_info': lambda: get_system_info(),
        'procesos': lambda: get_procesos(),
        'matar_proceso': lambda: matar_proceso(p.get('nombre', '')),
        'listar_carpeta': lambda: listar_carpeta(p.get('ruta')),
        'buscar_archivo': lambda: buscar_archivo(p.get('nombre', ''), p.get('donde')),
        'crear_carpeta': lambda: crear_carpeta(p.get('ruta', '')),
        'mover_archivo': lambda: mover_archivo(p.get('origen', ''), p.get('destino', '')),
        'copiar_archivo': lambda: copiar_archivo(p.get('origen', ''), p.get('destino', '')),
        'renombrar': lambda: renombrar_archivo(p.get('origen', ''), p.get('nombre', '')),
        'abrir_archivo': lambda: abrir_archivo(p.get('ruta', '')),
        'abrir_carpeta': lambda: abrir_carpeta(p.get('ruta')),
        'listar_ventanas': lambda: listar_ventanas(),
        'enfocar_ventana': lambda: enfocar_ventana(p.get('titulo', '')),
        'cerrar_ventana': lambda: cerrar_ventana_titulo(p.get('titulo', '')),
        'minimizar_ventana': lambda: minimizar_ventana(p.get('titulo')),
        'maximizar_ventana': lambda: maximizar_ventana(p.get('titulo')),
        'abrir_programa': lambda: abrir_programa(p.get('nombre', '')),
        'abrir_web': lambda: abrir_web(p.get('url', '')),
        'buscar_google': lambda: buscar_google(p.get('query', '')),
        'escribir': lambda: escribir_texto(p.get('texto', ''), float(p.get('delay', 0.03))),
        'tecla': lambda: tecla(p.get('key', '')),
        'hotkey': lambda: hotkey(*p.get('keys', [])),
        'portapapeles': lambda: portapapeles(p.get('texto', '')),
        'click': lambda: (pyautogui.click(int(p['x']), int(p['y'])) if 'x' in p else pyautogui.click()) or 'Click',
        'scroll': lambda: pyautogui.scroll(int(p.get('cantidad', 3))) or 'Scroll',
        'screenshot': lambda: screenshot(p.get('nombre', 'captura')),
        'volumen': lambda: volumen(p.get('accion', 'subir')),
        'notificar_windows': lambda: (notificar_windows(p.get('titulo', 'NOVA'), p.get('mensaje', '')), 'Notificación enviada')[1],
        'crear_word': lambda: crear_word(p.get('nombre', 'doc'), p.get('contenido', '')),
        'crear_pdf': lambda: crear_pdf(p.get('nombre', 'doc'), p.get('contenido', '')),
        'crear_imagen': lambda: crear_imagen(p.get('nombre', 'img'), p.get('texto', ''), p.get('ancho', 1920), p.get('alto', 1080)),
        'ejecutar_cmd': lambda: ejecutar_cmd(p.get('cmd', '')),
        'crear_auto': lambda: crear_automatizacion(p.get('nombre',''), p.get('cond_tipo',''), p.get('cond_valor',''), p.get('acc_tipo',''), p.get('acc_valor','')),
        'listar_autos': lambda: listar_automatizaciones(),
        'eliminar_auto': lambda: eliminar_automatizacion(p.get('nombre', '')),
    }
    fn = mapa.get(acc)
    if not fn:
        return jsonify({'ok': False, 'error': f'Acción desconocida: {acc}'}), 400
    try:
        resultado = fn()
        return jsonify({'ok': True, 'resultado': resultado})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/vision', methods=['POST'])
def vision_proxy():
    try:
        data = request.get_json() or {}
        nvidia_key = data.get('key', '')
        messages = data.get('messages', [])
        model = data.get('model', 'nvidia/nemotron-3-nano-omni-30b-a3b-reasoning')
        max_tokens = data.get('max_tokens', 1024)
        reasoning_budget = data.get('reasoning_budget', 512)

        payload = {'model': model, 'max_tokens': max_tokens, 'temperature': 0.4, 'messages': messages}
        if reasoning_budget:
            payload['reasoning_budget'] = reasoning_budget

        import requests as req_lib
        r = req_lib.post(
            'https://integrate.api.nvidia.com/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {nvidia_key}',
                'Content-Type': 'application/json',
                'Accept': 'application/json'
            },
            json=payload,
            timeout=30
        )
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

CODE_EXTENSIONS = {
    '.js', '.jsx', '.ts', '.tsx', '.swift', '.m', '.mm', '.h',
    '.kt', '.kts', '.java', '.dart', '.py', '.json', '.yaml', '.yml',
    '.plist', '.xcconfig', '.gradle', '.md', '.css', '.scss', '.html',
    '.vue', '.rb', '.go', '.rs', '.c', '.cpp', '.cs'
}

IGNORE_DIRS_PROYECTO = {
    'node_modules', '.git', '.expo', '.next', 'dist', 'build', 'Pods',
    'DerivedData', '.build', 'Carthage', '.dart_tool', '.gradle', '.idea',
    '.vscode', '__pycache__', '.venv', 'venv', 'coverage', '.cache',
    'xcuserdata', 'Vendor', 'vendor', '.xcworkspace'
}

IGNORE_FILES_PROYECTO = {
    'package-lock.json', 'yarn.lock', 'Podfile.lock', 'Gemfile.lock',
    'pubspec.lock', '.DS_Store'
}

MAX_ARCHIVOS_ESCANEO = 400
MAX_TAMANO_LECTURA_PROYECTO = 500_000

def _validar_ruta_proyecto(proyecto_root, ruta_relativa=None):
    try:
        root = Path(proyecto_root).resolve()
        if not root.exists() or not root.is_dir():
            return None, None, f'La ruta del proyecto no existe o no es una carpeta: {proyecto_root}'
        if ruta_relativa:
            ruta = (root / ruta_relativa).resolve()
            try:
                ruta.relative_to(root)
            except ValueError:
                return None, None, 'Ruta fuera del proyecto no permitida.'
            return root, ruta, None
        return root, None, None
    except Exception as e:
        return None, None, str(e)

@app.route('/api/proyecto/escanear', methods=['POST'])
def proyecto_escanear():
    try:
        data = request.get_json() or {}
        proyecto_root = data.get('ruta', '')
        root, _, err = _validar_ruta_proyecto(proyecto_root)
        if err:
            return jsonify({'ok': False, 'error': err}), 400

        archivos = []
        for p in root.rglob('*'):
            if len(archivos) >= MAX_ARCHIVOS_ESCANEO:
                break
            if p.is_dir():
                continue
            try:
                partes = p.relative_to(root).parts
            except ValueError:
                continue
            if any(part in IGNORE_DIRS_PROYECTO for part in partes):
                continue
            if p.name in IGNORE_FILES_PROYECTO:
                continue
            if p.suffix.lower() not in CODE_EXTENSIONS:
                continue
            try:
                tam = p.stat().st_size
            except Exception:
                tam = 0
            archivos.append({
                'ruta': str(p.relative_to(root)).replace('\\', '/'),
                'tamano': tam,
            })

        return jsonify({'ok': True, 'raiz': str(root), 'total': len(archivos), 'archivos': archivos})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/proyecto/leer', methods=['POST'])
def proyecto_leer():
    try:
        data = request.get_json() or {}
        proyecto_root = data.get('ruta_proyecto', '')
        archivo = data.get('archivo', '')
        root, ruta, err = _validar_ruta_proyecto(proyecto_root, archivo)
        if err:
            return jsonify({'ok': False, 'error': err}), 400
        if not ruta.exists() or not ruta.is_file():
            return jsonify({'ok': False, 'error': f'Archivo no encontrado: {archivo}'}), 404

        tam = ruta.stat().st_size
        if tam > MAX_TAMANO_LECTURA_PROYECTO:
            return jsonify({'ok': False, 'error': f'Archivo demasiado grande ({tam} bytes).'}), 400

        try:
            contenido = ruta.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            return jsonify({'ok': False, 'error': 'El archivo no es texto UTF-8 (¿binario?).'}), 400

        return jsonify({'ok': True, 'archivo': archivo, 'contenido': contenido})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/nvidia_chat', methods=['POST'])
def nvidia_chat_proxy():
    try:
        data = request.get_json() or {}
        nvidia_key = data.get('key', '')
        messages = data.get('messages', [])
        model = data.get('model', 'qwen/qwen3-coder-480b-a35b-instruct')
        max_tokens = data.get('max_tokens', 2048)

        import requests as req_lib
        r = req_lib.post(
            'https://integrate.api.nvidia.com/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {nvidia_key}',
                'Content-Type': 'application/json',
                'Accept': 'application/json'
            },
            json={'model': model, 'max_tokens': max_tokens, 'temperature': 0.3, 'messages': messages},
            timeout=45
        )
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("🤖 NOVA Agente v2 arrancando...")
    print(f"📁 Proyecto en: {NOVA_ROOT}")
    print(f"💾 Backups en: {BACKUPS_DIR}")
    print(f"🔔 Motor de notificaciones: {_notif_engine or 'NINGUNO'}")
    threading.Thread(target=monitor_proactivo, daemon=True).start()
    threading.Thread(target=monitor_automatizaciones, daemon=True).start()
    print("✅ Monitores activos")
    print("🚀 Servidor en http://localhost:4000")
    print("━" * 50)
    print("Capacidades:")
    print(" • Monitorización CPU/RAM/batería en tiempo real")
    print(" • Gestión de archivos y carpetas")
    print(" • Control de ventanas y aplicaciones")
    print(" • Automatizaciones por condición")
    print(" • Alertas proactivas")
    print(" • Proxy de noticias (Google News RSS)")
    print(" • Proxy de visión y código NVIDIA")
    print(" • Análisis de proyectos externos (solo lectura)")
    print(" • 🧬 AUTO-MEJORA con backups automáticos")
    print("━" * 50)
    app.run(host='127.0.0.1', port=4000, threaded=True)